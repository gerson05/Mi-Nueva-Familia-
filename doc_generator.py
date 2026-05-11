from docx import Document
from docx.shared import Inches

def generate_receipt(template_path: str, output_path: str, context: dict, image_path: str = None):
    # ...
    doc = Document(template_path)
    
    # Suponiendo que la información está en la primera tabla
    if not doc.tables:
        raise ValueError("El documento Word no contiene ninguna tabla.")
        
    table = doc.tables[0]
    
    # Iteramos sobre las filas y reemplazamos los valores en la segunda columna (index 1)
    # según el texto de la primera columna (index 0)
    for row in table.rows:
        if len(row.cells) >= 2:
            key_cell = row.cells[0].text.strip().lower()
            val_cell = row.cells[1]
            
            # Limpiamos el texto original de la segunda celda y agregamos el nuevo con un espacio inicial
            if "foro" in key_cell:
                val_cell.text = " " + context.get("foro", "")
            elif "nombre patrocinador" in key_cell:
                val_cell.text = " " + context.get("nombre", "")
            elif "número de identificación" in key_cell or "numero de identificacion" in key_cell:
                val_cell.text = " " + context.get("identificacion", "")
            elif "teléfono" in key_cell or "telefono" in key_cell:
                val_cell.text = " " + context.get("telefono", "")
            elif "correo electrónico" in key_cell or "correo electronico" in key_cell:
                val_cell.text = " " + context.get("correo", "")
            elif "valor consignación" in key_cell or "valor consignacion" in key_cell:
                val_cell.text = " " + context.get("valor", "")
            elif "fecha de consignación" in key_cell or "fecha de consignacion" in key_cell:
                val_cell.text = " " + context.get("fecha", "")
            elif "número de comprobante" in key_cell or "numero de comprobante" in key_cell:
                val_cell.text = " " + context.get("comprobante", "")
            elif "mes del aporte" in key_cell:
                val_cell.text = " " + context.get("mes", "")
            elif "niños patrocinados" in key_cell or "ninos patrocinados" in key_cell:
                val_cell.text = " " + context.get("ninos", "")
            elif "año de patrocinio" in key_cell or "ano de patrocinio" in key_cell:
                val_cell.text = " " + context.get("ano_patrocinio", "")
                
    if image_path:
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Crear un nuevo párrafo antes de la tabla para la imagen
        new_p = OxmlElement('w:p')
        table._element.addprevious(new_p)
        
        # Crear un párrafo vacío adicional para que sirva de espacio/margen inferior
        spacer_p = OxmlElement('w:p')
        table._element.addprevious(spacer_p)
        
        spacer_p2 = OxmlElement('w:p')
        table._element.addprevious(spacer_p2)
        
        p = Paragraph(new_p, doc._body)
        if context.get("alineacion") == "Central":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        try:
            # Añadir la imagen con un alto restringido para asegurar que quepa en 1 sola página
            run = p.add_run()
            run.add_picture(image_path, height=Inches(3.5))
        except Exception as e:
            print("No se pudo añadir la imagen al word:", e)

    doc.save(output_path)
